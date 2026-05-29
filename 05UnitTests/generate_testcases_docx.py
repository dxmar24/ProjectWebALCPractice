from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "TestCases_AmericanLatinClass.md"
TARGET = ROOT / "TestCases_AmericanLatinClass.docx"


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def set_cell_width(cell, width_inches: float) -> None:
    width = int(width_inches * 1440)
    properties = cell._tc.get_or_add_tcPr()
    tc_width = properties.first_child_found_in("w:tcW")

    if tc_width is None:
        tc_width = OxmlElement("w:tcW")
        properties.append(tc_width)

    tc_width.set(qn("w:w"), str(width))
    tc_width.set(qn("w:type"), "dxa")


def set_cell_text(cell, text: str, *, bold: bool = False, size: float = 7.7) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text.replace("`", ""))
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)


def set_document_defaults(document: Document) -> None:
    section = document.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9)
    styles["Normal"].paragraph_format.space_after = Pt(4)
    styles["Normal"].paragraph_format.line_spacing = 1.0

    for style_name, size in [("Heading 1", 14), ("Heading 2", 11), ("Heading 3", 10)]:
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.bold = True
        style.font.size = Pt(size)
        style.paragraph_format.space_before = Pt(6)
        style.paragraph_format.space_after = Pt(4)


def add_paragraph(document: Document, text: str, *, align=None, bold: bool = False, size: float = 9) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.0

    if align is not None:
        paragraph.alignment = align

    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)


def add_heading(document: Document, text: str, level: int) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = f"Heading {min(level, 3)}"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(14 if level == 1 else 11)
    run.bold = True


def split_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def is_separator(line: str) -> bool:
    parts = split_table_row(line)
    return bool(parts) and all(set(part.replace(":", "").strip()) <= {"-"} and "-" in part for part in parts)


def table_widths(headers: list[str]) -> list[float]:
    if headers == ["Feature ID", "Feature", "Total Test Cases", "Functional", "Pending / Review"]:
        return [0.85, 3.25, 1.0, 0.85, 1.2]

    return [0.82, 1.25, 1.95, 2.25, 0.85]


def add_table(document: Document, rows: list[list[str]]) -> None:
    widths = table_widths(rows[0])
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False

    for row_index, row in enumerate(rows):
        for col_index, text in enumerate(row):
            cell = table.cell(row_index, col_index)
            set_cell_width(cell, widths[col_index])
            set_cell_text(cell, text, bold=row_index == 0, size=7.4 if len(rows[0]) == 5 else 8)

            if row_index == 0:
                set_cell_shading(cell, "D9EAF7")
            elif text.lower() in {"functional", "pending", "review"}:
                set_cell_shading(cell, "EAF5EA" if text.lower() == "functional" else "FFF2CC")

    document.add_paragraph()


def build_document() -> None:
    document = Document()
    set_document_defaults(document)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    index = 0

    while index < len(lines):
        raw = lines[index]
        line = raw.strip()

        if line == "":
            index += 1
            continue

        if line == r"\newpage":
            document.add_page_break()
            index += 1
            continue

        if line.startswith("|") and index + 1 < len(lines) and is_separator(lines[index + 1]):
            rows = [split_table_row(line)]
            index += 2

            while index < len(lines) and lines[index].strip().startswith("|"):
                rows.append(split_table_row(lines[index]))
                index += 1

            add_table(document, rows)
            continue

        if line.startswith("# "):
            add_heading(document, line[2:].strip(), 1)
            index += 1
            continue

        if line.startswith("## "):
            add_heading(document, line[3:].strip(), 2)
            index += 1
            continue

        if line.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.space_after = Pt(2)
            run = paragraph.add_run(line[2:].replace("`", ""))
            run.font.name = "Arial"
            run.font.size = Pt(9)
            index += 1
            continue

        add_paragraph(document, line.replace("**", "").replace("`", ""), size=9)
        index += 1

    document.save(TARGET)


if __name__ == "__main__":
    build_document()
